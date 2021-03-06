# -*- coding: utf-8 -*-

from PyQt5 import QtCore, QtGui, QtWidgets
from docx import Document #pip install python-docx
import mysql.connector
import time

cnx = mysql.connector.connect(user='root', password='i130813',
                              host='127.0.0.1',
                              database='mydb')
cursor = cnx.cursor()

trip_data = {}


class Ui_MainWindow(object):
    def setupUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.showFullScreen()
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 1, 0, 1, 1)
        self.fromedit = QtWidgets.QLineEdit(self.centralwidget)
        self.fromedit.setObjectName("fromedit")
        self.gridLayout.addWidget(self.fromedit, 1, 1, 1, 1)
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.gridLayout.addWidget(self.label_3, 2, 0, 1, 1)
        self.whereedit = QtWidgets.QLineEdit(self.centralwidget)
        self.whereedit.setObjectName("whereedit")
        self.gridLayout.addWidget(self.whereedit, 2, 1, 1, 1)
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.gridLayout.addWidget(self.label_4, 3, 0, 1, 1)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_5.setFont(font)
        self.label_5.setObjectName("label_5")
        self.gridLayout.addWidget(self.label_5, 4, 0, 1, 1)
        self.searchbutton = QtWidgets.QPushButton(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(12)
        self.searchbutton.setFont(font)
        self.searchbutton.setCursor(QtGui.QCursor(QtCore.Qt.OpenHandCursor))
        self.searchbutton.setObjectName("searchbutton")
        self.gridLayout.addWidget(self.searchbutton, 5, 1, 1, 1)
        self.label = QtWidgets.QLabel(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(40)
        self.label.setFont(font)
        self.label.setTextFormat(QtCore.Qt.AutoText)
        self.label.setAlignment(QtCore.Qt.AlignCenter)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 0, 0, 1, 2)
        self.whenedit = QtWidgets.QDateEdit(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.whenedit.setFont(font)
        self.whenedit.setCursor(QtGui.QCursor(QtCore.Qt.ArrowCursor))
        self.whenedit.setDateTime(QtCore.QDateTime(QtCore.QDate(2019, 3, 1), QtCore.QTime(0, 0, 0)))
        self.whenedit.setCalendarPopup(True)
        self.whenedit.setObjectName("whenedit")
        self.gridLayout.addWidget(self.whenedit, 3, 1, 1, 1)
        self.transportedit = QtWidgets.QComboBox(self.centralwidget)
        self.transportedit.setCursor(QtGui.QCursor(QtCore.Qt.OpenHandCursor))
        self.transportedit.setObjectName("transportedit")
        self.gridLayout.addWidget(self.transportedit, 4, 1, 1, 1)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 549, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label_2.setText(_translate("MainWindow", "Откуда"))
        self.label_3.setText(_translate("MainWindow", "Куда"))
        self.label_4.setText(_translate("MainWindow", "Когда"))
        self.label_5.setText(_translate("MainWindow", "На чем"))
        self.searchbutton.setText(_translate("MainWindow", "Поиск"))
        self.label.setText(_translate("MainWindow", "Билеты на транспорт"))

        self.transportedit.addItem("Автобус")
        self.transportedit.addItem("Поезд")
        self.transportedit.addItem("Самолет")

        self.searchbutton.clicked.connect(self.prepareGlobalViewUi)

    def prepareGlobalViewUi(self):
        if self.whenedit.text() != "" and self.fromedit.text() != "" and self.whereedit.text() != "" and self.transportedit.itemText(self.transportedit.currentIndex()) != "":
            trip_data.update({"where": self.whereedit.text()})
            trip_data.update({"from": self.fromedit.text()})
            data = self.whenedit.text().split(".")
            data = str(data[2]) + "-" + str(data[1]) + "-" + str(data[0])
            trip_data.update({"when": data})
            trip_data.update({"transport": self.transportedit.itemText(self.transportedit.currentIndex())})
            self.setupGlobalViewUi()
        else:
            if self.whenedit.text() != "":
                pass

    def setupGlobalViewUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.showFullScreen()
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.routenumberlabel = QtWidgets.QLabel(self.centralwidget)
        self.routenumberlabel.setObjectName("routenumberlabel")
        self.horizontalLayout.addWidget(self.routenumberlabel)
        self.fromlabel = QtWidgets.QLabel(self.centralwidget)
        self.fromlabel.setObjectName("fromlabel")
        self.horizontalLayout.addWidget(self.fromlabel)
        self.departtimelabel = QtWidgets.QLabel(self.centralwidget)
        self.departtimelabel.setObjectName("departtimelabel")
        self.horizontalLayout.addWidget(self.departtimelabel)
        self.timeinttravellabel = QtWidgets.QLabel(self.centralwidget)
        self.timeinttravellabel.setObjectName("timeinttravellabel")
        self.horizontalLayout.addWidget(self.timeinttravellabel)
        self.wherelabel = QtWidgets.QLabel(self.centralwidget)
        self.wherelabel.setObjectName("wherelabel")
        self.horizontalLayout.addWidget(self.wherelabel)
        self.arrivetimelabel = QtWidgets.QLabel(self.centralwidget)
        self.arrivetimelabel.setObjectName("arrivetimelabel")
        self.horizontalLayout.addWidget(self.arrivetimelabel)
        self.costlabel = QtWidgets.QLabel(self.centralwidget)
        self.costlabel.setObjectName("costlabel")
        self.horizontalLayout.addWidget(self.costlabel)
        self.filler = QtWidgets.QLabel(self.centralwidget)
        self.filler.setObjectName("filler")
        self.horizontalLayout.addWidget(self.filler)
        self.verticalLayout.addLayout(self.horizontalLayout)
        self.scrollArea = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 836, 541))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.gridLayoutWidget = QtWidgets.QWidget(self.scrollAreaWidgetContents)
        self.gridLayoutWidget.setGeometry(QtCore.QRect(0, 0, 841, 571))
        self.gridLayoutWidget.setObjectName("gridLayoutWidget")
        self.gridLayout = QtWidgets.QGridLayout(self.gridLayoutWidget)
        self.gridLayout.setContentsMargins(0, 0, 0, 0)
        self.gridLayout.setObjectName("gridLayout")
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.verticalLayout.addWidget(self.scrollArea)
        self.backtomainbutton = QtWidgets.QPushButton(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.backtomainbutton.sizePolicy().hasHeightForWidth())
        self.backtomainbutton.setSizePolicy(sizePolicy)
        self.backtomainbutton.setObjectName("pushButton")
        self.verticalLayout.addWidget(self.backtomainbutton)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 856, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateGlobalViewUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateGlobalViewUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.routenumberlabel.setText(_translate("MainWindow", "Рейс"))
        self.fromlabel.setText(_translate("MainWindow", "Откуда"))
        self.departtimelabel.setText(_translate("MainWindow", "Время отправления"))
        self.timeinttravellabel.setText(_translate("MainWindow", "Время в пути"))
        self.wherelabel.setText(_translate("MainWindow", "Куда"))
        self.arrivetimelabel.setText(_translate("MainWindow", "Время прибытия"))
        self.costlabel.setText(_translate("MainWindow", "Цена"))
        self.filler.setText(_translate("MainWindow", " "))
        self.backtomainbutton.setText(_translate("MainWindow", "Назад"))
        self.backtomainbutton.clicked.connect(self.setupUi)

        if trip_data.get("from") != None:
            query = "select * from trip where FromCity= %s and ToCity= %s and DateDeparture= %s and transport = %s;"
            data = (trip_data.get("from"), trip_data.get("where"), trip_data.get("when"), trip_data.get("transport"))
            cursor.execute(query, data)

        j = 0
        k = 0

        x = 0
        y = 0
        font = QtGui.QFont()
        font.setPointSize(10)

        for item in cursor:
            item_group = QtWidgets.QGroupBox("Рейс: " + str(item[0]))
            categorieslayout = QtWidgets.QVBoxLayout(item_group)
            self.gridLayout.addWidget(item_group, x, y, 1, 1)
            item_group.setFont(font)
            for value in item:
                if k == 0:
                    k += 1
                    j = str(value)
                    trip_data.update({j: str(value)})
                    continue
                line_item = QtWidgets.QLabel(str(value))
                categorieslayout.addWidget(line_item)
                k += 1
                if k == 11:
                    line_item = QtWidgets.QPushButton("Купить")
                    categorieslayout.addWidget(line_item)
                    line_item.clicked.connect(lambda state, row=j: open_info(row))
                    k = 0
            if x == 4:
                y += 1
                x = 0
            x += 1

        def open_info(row):
            id = trip_data.get(row)
            trip_data.clear()
            trip_data.update({"id": id})
            self.setupRouteInfoUi()

    def setupRouteInfoUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.showFullScreen()
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.surnameedit = QtWidgets.QLineEdit(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Maximum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.surnameedit.sizePolicy().hasHeightForWidth())
        self.surnameedit.setSizePolicy(sizePolicy)
        self.surnameedit.setObjectName("surnameedit")
        self.gridLayout.addWidget(self.surnameedit, 3, 1, 1, 1)
        self.nextButton = QtWidgets.QPushButton(self.centralwidget)
        self.nextButton.setObjectName("nextButton")
        self.gridLayout.addWidget(self.nextButton, 11, 0, 1, 2)
        self.passedit = QtWidgets.QLineEdit(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Maximum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.passedit.sizePolicy().hasHeightForWidth())
        self.passedit.setSizePolicy(sizePolicy)
        self.passedit.setObjectName("passedit")
        self.gridLayout.addWidget(self.passedit, 7, 1, 1, 1)
        self.passlabel = QtWidgets.QLabel(self.centralwidget)
        self.passlabel.setObjectName("passlabel")
        self.gridLayout.addWidget(self.passlabel, 6, 1, 1, 1)
        self.nameedit = QtWidgets.QLineEdit(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Maximum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.nameedit.sizePolicy().hasHeightForWidth())
        self.nameedit.setSizePolicy(sizePolicy)
        self.nameedit.setObjectName("nameedit")
        self.gridLayout.addWidget(self.nameedit, 1, 1, 1, 1)
        self.namelabel = QtWidgets.QLabel(self.centralwidget)
        self.namelabel.setObjectName("namelabel")
        self.gridLayout.addWidget(self.namelabel, 0, 1, 1, 1)
        self.fromlabel = QtWidgets.QLabel(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Minimum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.fromlabel.sizePolicy().hasHeightForWidth())
        self.fromlabel.setSizePolicy(sizePolicy)
        self.fromlabel.setObjectName("fromlabel")
        self.gridLayout.addWidget(self.fromlabel, 2, 0, 1, 1)
        self.surnamelabel = QtWidgets.QLabel(self.centralwidget)
        self.surnamelabel.setObjectName("surnamelabel")
        self.gridLayout.addWidget(self.surnamelabel, 2, 1, 1, 1)
        self.wherelabel = QtWidgets.QLabel(self.centralwidget)
        self.wherelabel.setObjectName("wherelabel")
        self.gridLayout.addWidget(self.wherelabel, 4, 0, 1, 1)
        self.patrlabel = QtWidgets.QLabel(self.centralwidget)
        self.patrlabel.setObjectName("patrlabel")
        self.gridLayout.addWidget(self.patrlabel, 4, 1, 1, 1)
        self.routelabel = QtWidgets.QLabel(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.routelabel.sizePolicy().hasHeightForWidth())
        self.routelabel.setSizePolicy(sizePolicy)
        self.routelabel.setObjectName("routelabel")
        self.gridLayout.addWidget(self.routelabel, 0, 0, 1, 1)
        self.partedit = QtWidgets.QLineEdit(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.partedit.sizePolicy().hasHeightForWidth())
        self.partedit.setSizePolicy(sizePolicy)
        self.partedit.setObjectName("partedit")
        self.gridLayout.addWidget(self.partedit, 5, 1, 1, 1)
        self.whenlabel = QtWidgets.QLabel(self.centralwidget)
        self.whenlabel.setObjectName("whenlabel")
        self.gridLayout.addWidget(self.whenlabel, 6, 0, 1, 1)
        self.costlabel = QtWidgets.QLabel(self.centralwidget)
        self.costlabel.setObjectName("costlabel")
        self.gridLayout.addWidget(self.costlabel, 10, 0, 1, 1)
        self.companylabel = QtWidgets.QLabel(self.centralwidget)
        self.companylabel.setObjectName("companylabel")
        self.gridLayout.addWidget(self.companylabel, 9, 0, 1, 1)
        self.arrivelabel = QtWidgets.QLabel(self.centralwidget)
        self.arrivelabel.setObjectName("arrivelabel")
        self.gridLayout.addWidget(self.arrivelabel, 8, 0, 1, 1)
        self.backButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.backButton_2.setObjectName("backButton_2")
        self.gridLayout.addWidget(self.backButton_2, 12, 0, 1, 2)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 814, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateRouteInfoUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateRouteInfoUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.nextButton.setText(_translate("MainWindow", "Продолжить"))
        self.passlabel.setText(_translate("MainWindow", "Паспорт"))
        self.namelabel.setText(_translate("MainWindow", "Имя"))
        self.fromlabel.setText(_translate("MainWindow", "Из:"))
        self.surnamelabel.setText(_translate("MainWindow", "Фамилия"))
        self.wherelabel.setText(_translate("MainWindow", "В:"))
        self.patrlabel.setText(_translate("MainWindow", "Отчество"))
        self.routelabel.setText(_translate("MainWindow", "Рейс:"))
        self.whenlabel.setText(_translate("MainWindow", "Время отправления:"))
        self.costlabel.setText(_translate("MainWindow", "Стоимость:"))
        self.companylabel.setText(_translate("MainWindow", "Компания:"))
        self.arrivelabel.setText(_translate("MainWindow", "Время прибытия:"))
        self.backButton_2.setText(_translate("MainWindow", "Назад"))
        self.nextButton.clicked.connect(self.createpassenger)

        query = "select * from trip where idTrip= %s;"
        data = (str(trip_data.get("id")))
        cursor.execute(query, (data,))

        j = 0
        for item in cursor:
            for value in item:
                if j == 0:
                    j += 1
                    continue

                if j == 1:
                    self.routelabel.setText("Номер рейса:" + str(value))
                if j == 2:
                    self.fromlabel.setText("Из:" + str(value))
                if j == 3:
                    self.wherelabel.setText("В:" + str(value))
                if j == 4:
                    self.whenlabel.setText("Время отправления:" + str(value))
                if j == 5:
                    self.arrivelabel.setText("Время прибытия:" + str(value))
                if j == 6:
                    self.companylabel.setText("Перевозчик:" + str(value))
                if j == 7:
                    self.costlabel.setText("Цена:" + str(value))
                if j == 8:
                    text = self.whenlabel.text()
                    self.whenlabel.setText(text + " " + str(value))
                if j == 9:
                    text = self.arrivelabel.text()
                    self.arrivelabel.setText(text + " " + str(value))

                j += 1

    def createpassenger(self):
        if self.nameedit.text() != "" and self.surnameedit.text() != "" and self.partedit.text() != "" and self.passedit.text() != "":
            data = (self.nameedit.text(), self.surnameedit.text(), self.partedit.text(), self.passedit.text())
            query = "insert into passenger(NameP,SurnameP,Patronymic,Passport) values(%s,%s,%s,%s);"
            cursor.execute(query, data)
            cnx.commit()
            trip_data.update({"pass": self.passedit.text()})
            self.setupPurchaseUi()

    def setupPurchaseUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.showFullScreen()
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.nextButton = QtWidgets.QPushButton(self.centralwidget)
        self.nextButton.setObjectName("nextButton")
        self.gridLayout.addWidget(self.nextButton, 8, 0, 1, 2)
        self.backButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.backButton_2.setObjectName("backButton_2")
        self.gridLayout.addWidget(self.backButton_2, 9, 0, 1, 2)
        self.passlabel = QtWidgets.QLabel(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Maximum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.passlabel.sizePolicy().hasHeightForWidth())
        self.passlabel.setSizePolicy(sizePolicy)
        self.passlabel.setObjectName("passlabel")
        self.gridLayout.addWidget(self.passlabel, 3, 1, 1, 1)
        self.namelabel = QtWidgets.QLabel(self.centralwidget)
        self.namelabel.setObjectName("namelabel")
        self.gridLayout.addWidget(self.namelabel, 0, 1, 1, 1)
        self.fromlabel = QtWidgets.QLabel(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Minimum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.fromlabel.sizePolicy().hasHeightForWidth())
        self.fromlabel.setSizePolicy(sizePolicy)
        self.fromlabel.setObjectName("fromlabel")
        self.gridLayout.addWidget(self.fromlabel, 1, 0, 1, 1)
        self.surnamelabel = QtWidgets.QLabel(self.centralwidget)
        self.surnamelabel.setObjectName("surnamelabel")
        self.gridLayout.addWidget(self.surnamelabel, 1, 1, 1, 1)
        self.wherelabel = QtWidgets.QLabel(self.centralwidget)
        self.wherelabel.setObjectName("wherelabel")
        self.gridLayout.addWidget(self.wherelabel, 2, 0, 1, 1)
        self.patrlabel = QtWidgets.QLabel(self.centralwidget)
        self.patrlabel.setObjectName("patrlabel")
        self.gridLayout.addWidget(self.patrlabel, 2, 1, 1, 1)
        self.routelabel = QtWidgets.QLabel(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.routelabel.sizePolicy().hasHeightForWidth())
        self.routelabel.setSizePolicy(sizePolicy)
        self.routelabel.setObjectName("routelabel")
        self.gridLayout.addWidget(self.routelabel, 0, 0, 1, 1)
        self.whenlabel = QtWidgets.QLabel(self.centralwidget)
        self.whenlabel.setObjectName("whenlabel")
        self.gridLayout.addWidget(self.whenlabel, 3, 0, 1, 1)
        self.costlabel = QtWidgets.QLabel(self.centralwidget)
        self.costlabel.setObjectName("costlabel")
        self.gridLayout.addWidget(self.costlabel, 6, 0, 1, 1)
        self.companylabel = QtWidgets.QLabel(self.centralwidget)
        self.companylabel.setObjectName("companylabel")
        self.gridLayout.addWidget(self.companylabel, 5, 0, 1, 1)
        self.arrivelabel = QtWidgets.QLabel(self.centralwidget)
        self.arrivelabel.setObjectName("arrivelabel")
        self.gridLayout.addWidget(self.arrivelabel, 4, 0, 1, 1)
        self.cashlabel = QtWidgets.QLabel(self.centralwidget)
        self.cashlabel.setObjectName("label")
        self.gridLayout.addWidget(self.cashlabel, 7, 0, 1, 1)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 814, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslatePurchaseUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslatePurchaseUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.nextButton.setText(_translate("MainWindow", "Оплатить"))
        self.backButton_2.setText(_translate("MainWindow", "Отмена"))
        self.passlabel.setText(_translate("MainWindow", "Паспорт:"))
        self.namelabel.setText(_translate("MainWindow", "Имя:"))
        self.fromlabel.setText(_translate("MainWindow", "Из:"))
        self.surnamelabel.setText(_translate("MainWindow", "Фамилия:"))
        self.wherelabel.setText(_translate("MainWindow", "В:"))
        self.patrlabel.setText(_translate("MainWindow", "Отчество:"))
        self.routelabel.setText(_translate("MainWindow", "Рейс:"))
        self.whenlabel.setText(_translate("MainWindow", "Время отправления:"))
        self.costlabel.setText(_translate("MainWindow", "Стоимость:"))
        self.companylabel.setText(_translate("MainWindow", "Компания:"))
        self.arrivelabel.setText(_translate("MainWindow", "Время прибытия:"))
        self.cashlabel.setText(_translate("MainWindow", "Внесено: "))
        self.nextButton.clicked.connect(self.pay)

        query = "select * from trip where idTrip= %s;"
        data = (str(trip_data.get("id")))
        cursor.execute(query, (data,))
        j = 0
        for item in cursor:
            for value in item:
                if j == 0:
                    j += 1
                    continue

                if j == 1:
                    self.routelabel.setText("Номер рейса:" + str(value))
                if j == 2:
                    self.fromlabel.setText("Из:" + str(value))
                if j == 3:
                    self.wherelabel.setText("В:" + str(value))
                if j == 4:
                    self.whenlabel.setText("Время отправления:" + str(value))
                if j == 5:
                    self.arrivelabel.setText("Время прибытия:" + str(value))
                if j == 6:
                    self.companylabel.setText("Перевозчик:" + str(value))
                if j == 7:
                    self.costlabel.setText("Цена:" + str(value))
                if j == 8:
                    text = self.whenlabel.text()
                    self.whenlabel.setText(text + " " + str(value))
                if j == 9:
                    text = self.arrivelabel.text()
                    self.arrivelabel.setText(text + " " + str(value))

                j += 1

        query = "select * from passenger where Passport= %s;"
        data = (str(trip_data.get("pass")))
        cursor.execute(query, (data,))
        j = 0
        for item in cursor:
            for value in item:
                if j == 0:
                    trip_data.update({"pass_id": value})
                    j += 1
                    continue

                if j == 1:
                    self.namelabel.setText("Имя: " + str(value))
                if j == 2:
                    self.surnamelabel.setText("Фамилия: " + str(value))
                if j == 3:
                    self.patrlabel.setText("Отчество: " + str(value))
                if j == 4:
                    self.passlabel.setText("Паспорт: " + str(value))

                j += 1

    def pay(self):
        document = Document()
        document.add_heading('Транспортная квитанция', 0)

        query = "select * from trip where idTrip= %s;"
        data = (str(trip_data.get("id")))
        cursor.execute(query, (data,))
        j = 0
        for item in cursor:
            for value in item:
                if j == 0:
                    j += 1
                    continue

                if j == 1:
                    p = document.add_paragraph("Рейс: " + str(value))
                if j == 2:
                    p.add_run("\nИз: " + str(value))
                if j == 3:
                    p.add_run("\nВ: " + str(value))
                if j == 4:
                    p.add_run("\nДата отправления: " + str(value))
                if j == 5:
                    p.add_run("\nДата прибытия: " + str(value))
                if j == 6:
                    p.add_run("\nПеревозчик: " + str(value))
                if j == 7:
                    p.add_run("\nЦена: " + str(value))
                if j == 8:
                    p.add_run("\nВремя отправления:" + str(value))
                if j == 9:
                    p.add_run("\nВремя прибытия: " + str(value))

                j += 1

        query = "select * from passenger where Passport= %s;"
        data = (str(trip_data.get("pass")))
        cursor.execute(query, (data,))
        j = 0
        for item in cursor:
            for value in item:
                if j == 0:
                    trip_data.update({"pass_id": value})
                    j += 1
                    continue

                if j == 1:
                    p.add_run("\nИмя: " + str(value))
                if j == 2:
                    p.add_run("\nФамилия: " + str(value))
                if j == 3:
                    p.add_run("\nОтчество: " + str(value))
                if j == 4:
                    p.add_run("\nПаспорт: " + str(value))

                j += 1

        document.add_page_break()
        document.save('Билет.docx')
        self.setupendUi()

    def setupendUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.showFullScreen()
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.label = QtWidgets.QLabel(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label.sizePolicy().hasHeightForWidth())
        self.label.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setPointSize(40)
        self.label.setFont(font)
        self.label.setLayoutDirection(QtCore.Qt.RightToLeft)
        self.label.setAlignment(QtCore.Qt.AlignCenter)
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateendUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateendUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label.setText(_translate("MainWindow", "Спасибо за покупку!"))
        time.sleep(1)
        self.setupUi()


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi()
    app.setStyle("Fusion")
    MainWindow.show()
    sys.exit(app.exec_())
